#!/usr/bin/env python3
import argparse
import json
import logging
import sys
from pathlib import Path
from typing import Dict, List

sys.path.insert(0, str(Path(__file__).resolve().parent))

from utils import setup_logging

logger = logging.getLogger(__name__)


def compute_metrics(
    true_positives: int,
    false_positives: int,
    false_negatives: int,
) -> Dict[str, float]:
    precision = _safe_division(true_positives, true_positives + false_positives)
    recall = _safe_division(true_positives, true_positives + false_negatives)
    f1 = _safe_division(2 * precision * recall, precision + recall)
    accuracy = _safe_division(
        true_positives, true_positives + false_positives + false_negatives
    )
    return {
        "precision": precision,
        "recall": recall,
        "f1_score": f1,
        "accuracy": accuracy,
        "true_positives": true_positives,
        "false_positives": false_positives,
        "false_negatives": false_negatives,
    }


def _safe_division(numerator: int, denominator: int) -> float:
    if denominator == 0:
        return 0.0
    return round((numerator / denominator) * 100, 1)


def build_confusion_matrix(metrics: Dict[str, float]) -> str:
    tp = int(metrics["true_positives"])
    fp = int(metrics["false_positives"])
    fn = int(metrics["false_negatives"])
    header = "|                | Predicted PII | Predicted Non-PII |"
    sep = "|----------------|---------------|-------------------|"
    tp_row = f"| Actual PII     | {tp:<13} | {fn:<17} |"
    fn_row = f"| Actual Non-PII | {fp:<13} | {'-':>17} |"
    return "\n".join([header, sep, tp_row, fn_row])


def normalize_detections(detections: List[Dict]) -> List[Dict]:
    normalized = []
    for entry in detections:
        if "text" in entry:
            normalized.append({"text": entry["text"]})
        elif "original" in entry:
            normalized.append({"text": entry["original"]})
    return normalized


def evaluate(
    detections: List[Dict],
    ground_truth: List[Dict],
    output_report: str = "evaluation_report.md",
) -> Dict[str, float]:
    det_normalized = normalize_detections(detections)

    gt_texts = {entry.get("text", "") for entry in ground_truth if entry.get("text")}
    det_texts = {entry.get("text", "") for entry in det_normalized if entry.get("text")}

    true_positives = len(gt_texts & det_texts)
    false_positives = len(det_texts - gt_texts)
    false_negatives = len(gt_texts - det_texts)

    metrics = compute_metrics(true_positives, false_positives, false_negatives)
    report = generate_report(metrics, true_positives, false_positives, false_negatives)
    save_report(report, output_report)
    generate_docx_report(report, "Evaluation_Report.docx")
    return metrics


def generate_report(
    metrics: Dict[str, float],
    tp: int,
    fp: int,
    fn: int,
) -> str:
    lines = []
    lines.append("# PII Redaction Evaluation Report")
    lines.append("")
    lines.append("## Evaluation Strategy")
    lines.append("")
    lines.append(
        "The evaluation uses **entity-level comparison** between the ground truth "
        "and the detected entities. Each PII entity is treated as a single unit: "
        "the full text string must match exactly for a detection to count as a "
        "true positive."
    )
    lines.append("")
    lines.append("### Why Entity-Level Evaluation?")
    lines.append("")
    lines.append(
        "Entity-level evaluation is preferred over token-level evaluation for PII "
        "detection because PII entities are semantically meaningful spans of text. "
        "Token-level comparison would fragment a name like Rahul Sharma into two "
        "tokens (Rahul and Sharma) and could count them independently, leading to "
        "misleading precision and recall scores. Entity-level evaluation ensures "
        "that each complete PII span is correctly identified as a whole, reflecting "
        "the real-world requirement that a detector either catches or misses an "
        "entire piece of sensitive information."
    )
    lines.append("")
    lines.append("### Key Metrics")
    lines.append("")
    lines.append("- **True Positives (TP)**: Entities correctly detected as PII.")
    lines.append("- **False Positives (FP)**: Non-PII entities incorrectly flagged as PII.")
    lines.append("- **False Negatives (FN)**: PII entities missed by the detector.")
    lines.append("")
    lines.append("## Metrics")
    lines.append("")
    lines.append("### Formulas")
    lines.append("")
    lines.append("```")
    lines.append("Precision = TP / (TP + FP)")
    lines.append("Recall    = TP / (TP + FN)")
    lines.append("F1        = 2 * Precision * Recall / (Precision + Recall)")
    lines.append("Accuracy  = TP / (TP + FP + FN)")
    lines.append("```")
    lines.append("")
    lines.append("### Results")
    lines.append("")
    lines.append(f"- **Precision**: {metrics['precision']}%")
    lines.append(f"- **Recall**: {metrics['recall']}%")
    lines.append(f"- **F1 Score**: {metrics['f1_score']}%")
    lines.append(f"- **Accuracy**: {metrics['accuracy']}%")
    lines.append("")
    lines.append("## Confusion Matrix")
    lines.append("")
    lines.append("```")
    lines.append(build_confusion_matrix(metrics))
    lines.append("```")
    lines.append("")
    lines.append("## Observations")
    lines.append("")
    lines.append(
        "- Regex-based detection provides excellent precision for structured PII "
        " formats such as email addresses, phone numbers, SSNs, credit card numbers, "
        " IP addresses, and dates of birth."
    )
    lines.append(
        "- Named Entity Recognition (NER) may occasionally miss uncommon names "
        "or organizations that do not appear in the training data."
    )
    lines.append(
        "- Credit card numbers are validated using the Luhn algorithm before being "
        " classified as PII, reducing false positives from random number sequences."
    )
    lines.append(
        "- Full name detection relies on capitalized word sequence heuristics, which "
        " can produce false positives when section headers or labels contain multiple "
        " capitalized words in sequence."
    )
    lines.append("")
    lines.append("## Limitations")
    lines.append("")
    lines.append(
        "- OCR documents are not evaluated; the tool processes native .docx files only."
    )
    lines.append(
        "- Unusual or non-standard address formats may not be detected by the "
        " regex-based fallback."
    )
    lines.append(
        "- Rare or uncommon organization names may require additional training data "
        " or fine-tuned NER models for reliable detection."
    )
    lines.append(
        "- Entity-level evaluation requires exact string matching, so minor formatting "
        " differences (e.g., extra whitespace, different punctuation) may cause "
        " false negatives."
    )
    lines.append("")
    return "\n".join(lines)


def generate_docx_report(report_md: str, output_path: str) -> None:
    from docx import Document
    doc = Document()
    for line in report_md.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("```"):
            continue
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info("DOCX report saved to %s", output_path)


def save_report(report: str, output_path: str) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(report)
    logger.info("Evaluation report saved to %s", output_path)


def load_ground_truth(path: str) -> List[Dict]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    logger.info("Loaded ground truth with %d entries", len(data))
    return data


def load_detections(mapping_path: str) -> List[Dict]:
    with open(mapping_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    detections = []
    for original, replacement in data.items():
        detections.append({"text": original})
    logger.info("Loaded %d detections from mapping", len(detections))
    return detections


def build_cli_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Evaluate PII redaction results against ground truth"
    )
    parser.add_argument(
        "--ground-truth",
        required=True,
        help="Path to the ground truth JSON file",
    )
    parser.add_argument(
        "--mapping",
        default="mapping.json",
        help="Path to the mapping.json file (default: mapping.json)",
    )
    parser.add_argument(
        "--output",
        default="evaluation_report.md",
        help="Path to save the evaluation report (default: evaluation_report.md)",
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR"],
        help="Logging level (default: INFO)",
    )
    return parser


def main() -> None:
    parser = build_cli_parser()
    args = parser.parse_args()
    setup_logging(args.log_level)

    ground_truth = load_ground_truth(args.ground_truth)
    detections = load_detections(args.mapping)

    print(f"Ground truth entries: {len(ground_truth)}")
    print(f"Detected entries: {len(detections)}")

    metrics = evaluate(detections, ground_truth, args.output)

    print(f"Precision : {metrics['precision']}%")
    print(f"Recall    : {metrics['recall']}%")
    print(f"F1        : {metrics['f1_score']}%")
    print(f"Accuracy  : {metrics['accuracy']}%")
    logger.info("Evaluation report saved to %s", args.output)


if __name__ == "__main__":
    main()