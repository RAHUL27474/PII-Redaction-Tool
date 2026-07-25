import logging
import json
from pathlib import Path
from typing import Any, Dict, List, Tuple

logger = logging.getLogger(__name__)


def setup_logging(level: str = "INFO") -> None:
    logging.basicConfig(
        level=getattr(logging, level.upper(), logging.INFO),
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )


def read_docx(docx_path: str) -> Any:
    from docx import Document
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Input file not found: {docx_path}")
    logger.info("Reading document: %s", docx_path)
    return Document(docx_path)


def write_docx(doc: Any, output_path: str) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    logger.info("Writing redacted document to: %s", output_path)
    doc.save(output_path)


def load_mapping(mapping_path: str) -> Dict[str, str]:
    path = Path(mapping_path)
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        logger.info("Loaded mapping with %d entries from %s", len(data), mapping_path)
        return data
    logger.info("No existing mapping found at %s, starting fresh", mapping_path)
    return {}


def save_mapping(mapping: Dict[str, str], mapping_path: str) -> None:
    path = Path(mapping_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(mapping, f, indent=2, ensure_ascii=False)
    logger.info("Saved mapping with %d entries to %s", len(mapping), mapping_path)


def replace_in_paragraph(para: Any, mapping: Dict[str, str]) -> bool:
    original_text = para.text
    replaced_text = original_text
    for original, replacement in mapping.items():
        replaced_text = replaced_text.replace(original, replacement)
    if replaced_text != original_text:
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = replaced_text
        else:
            para.add_run(replaced_text)
        return True
    return False


def replace_in_table(table: Any, mapping: Dict[str, str]) -> int:
    count = 0
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if replace_in_paragraph(para, mapping):
                    count += 1
    return count


def apply_replacements(doc: Any, mapping: Dict[str, str]) -> int:
    total = 0
    for table in doc.tables:
        total += replace_in_table(table, mapping)
    for para in doc.paragraphs:
        if replace_in_paragraph(para, mapping):
            total += 1
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for para in header.paragraphs:
                    if replace_in_paragraph(para, mapping):
                        total += 1
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for para in footer.paragraphs:
                    if replace_in_paragraph(para, mapping):
                        total += 1
    logger.info("Applied %d paragraph replacements", total)
    return total